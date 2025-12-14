"""
Document parsers for various file formats
Supports PDF, DOCX, HTML, and plain text files
"""

import os
import re
from pathlib import Path
from typing import Tuple, Optional

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    import requests
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False


def parse_pdf(path: str) -> str:
    """
    Parse PDF file and extract text content.
    
    Args:
        path: Path to PDF file
        
    Returns:
        Extracted text content as string
        
    Raises:
        FileNotFoundError: If file doesn't exist
        ValueError: If PDF parsing fails or pdfplumber not available
    """
    if not PDFPLUMBER_AVAILABLE:
        raise ValueError("pdfplumber is required for PDF parsing. Install with: pip install pdfplumber")
    
    if not os.path.exists(path):
        raise FileNotFoundError(f"PDF file not found: {path}")
    
    text_parts = []
    
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
    except Exception as e:
        raise ValueError(f"Failed to parse PDF {path}: {str(e)}")
    
    return "\n\n".join(text_parts)


def parse_docx(path: str) -> str:
    """
    Parse DOCX file and extract text content.
    
    Args:
        path: Path to DOCX file
        
    Returns:
        Extracted text content as string
        
    Raises:
        FileNotFoundError: If file doesn't exist
        ValueError: If DOCX parsing fails or python-docx not available
    """
    if not DOCX_AVAILABLE:
        raise ValueError("python-docx is required for DOCX parsing. Install with: pip install python-docx")
    
    if not os.path.exists(path):
        raise FileNotFoundError(f"DOCX file not found: {path}")
    
    try:
        doc = DocxDocument(path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX {path}: {str(e)}")


def parse_html(path_or_url: str) -> str:
    """
    Parse HTML file or URL and extract text content.
    Removes script and style tags, extracts main content.
    
    Args:
        path_or_url: Path to HTML file or URL
        
    Returns:
        Extracted text content as string
        
    Raises:
        FileNotFoundError: If file doesn't exist
        ValueError: If HTML parsing fails or beautifulsoup4 not available
    """
    if not HTML_AVAILABLE:
        raise ValueError("beautifulsoup4 and requests are required for HTML parsing")
    
    try:
        # Check if it's a URL or file path
        if path_or_url.startswith(('http://', 'https://')):
            response = requests.get(path_or_url, timeout=10)
            response.raise_for_status()
            html_content = response.text
        else:
            if not os.path.exists(path_or_url):
                raise FileNotFoundError(f"HTML file not found: {path_or_url}")
            with open(path_or_url, 'r', encoding='utf-8') as f:
                html_content = f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Extract text
        text = soup.get_text(separator='\n')
        
        # Clean up excessive whitespace
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        return "\n".join(lines)
    except Exception as e:
        raise ValueError(f"Failed to parse HTML {path_or_url}: {str(e)}")


def parse_text(path: str) -> str:
    """
    Parse plain text file.
    
    Args:
        path: Path to text file
        
    Returns:
        File content as string
        
    Raises:
        FileNotFoundError: If file doesn't exist
        UnicodeDecodeError: If file encoding is not supported
    """
    if not os.path.exists(path):
        raise FileNotFoundError(f"Text file not found: {path}")
    
    # Try UTF-8 first, then fall back to latin-1
    encodings = ['utf-8', 'latin-1', 'cp1252']
    
    for encoding in encodings:
        try:
            with open(path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    
    raise ValueError(f"Could not decode text file {path} with any supported encoding")


def parse_file(path: str) -> Tuple[str, str]:
    """
    Automatically detect file type and parse accordingly.
    
    Args:
        path: Path to file
        
    Returns:
        Tuple of (text_content, detected_type)
        detected_type can be: 'pdf', 'docx', 'html', 'text', or 'unknown'
        
    Raises:
        FileNotFoundError: If file doesn't exist
        ValueError: If file type is unsupported or parsing fails
    """
    if not os.path.exists(path):
        raise FileNotFoundError(f"File not found: {path}")
    
    path_obj = Path(path)
    extension = path_obj.suffix.lower()
    
    # Map extensions to parser functions
    parsers = {
        '.pdf': (parse_pdf, 'pdf'),
        '.docx': (parse_docx, 'docx'),
        '.doc': (parse_docx, 'docx'),  # Try docx parser for .doc files
        '.html': (parse_html, 'html'),
        '.htm': (parse_html, 'html'),
        '.txt': (parse_text, 'text'),
        '.md': (parse_text, 'text'),
    }
    
    if extension in parsers:
        parser_func, doc_type = parsers[extension]
        try:
            text = parser_func(path)
            return text, doc_type
        except Exception as e:
            raise ValueError(f"Failed to parse {path} as {doc_type}: {str(e)}")
    
    # Try text parser as fallback
    try:
        text = parse_text(path)
        return text, 'text'
    except:
        raise ValueError(f"Unsupported file type: {extension} for file {path}")
